# backend/app/main.py
from fastapi import FastAPI, HTTPException, Request, UploadFile, File, Form, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Dict
import os
import re
import io
import numpy as np
from datetime import datetime, timedelta
from supabase import create_client, Client
from groq import Groq

from app.services.hf_model_loader import hf_loader

# ============================================
# ENVIRONMENT VARIABLES
# ============================================
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_ANON_KEY = os.environ.get("SUPABASE_ANON_KEY")
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")

print(f"🔐 Supabase configured: {SUPABASE_URL is not None}")
print(f"🤖 Groq configured: {GROQ_API_KEY is not None}")

# Initialize clients
supabase: Client = create_client(SUPABASE_URL, SUPABASE_ANON_KEY) if SUPABASE_URL and SUPABASE_ANON_KEY else None
groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# ============================================
# FASTAPI APP
# ============================================
app = FastAPI(title="ProphetLedger API", version="1.0.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://prophetledger.vercel.app",
        "https://prophet-ledger.vercel.app",
        "http://localhost:3000",
        "http://localhost:3001"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================
# AUTH HELPER FUNCTION
# ============================================
def get_current_user(request: Request):
    """Get current authenticated user"""
    auth_header = request.headers.get("Authorization")
    if not auth_header:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    token = auth_header.replace("Bearer ", "")
    if not supabase:
        raise HTTPException(status_code=500, detail="Supabase not configured")
    
    try:
        user = supabase.auth.get_user(token)
        return user.user
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid token")


# ============================================
# DEPENDENCIES FOR FILE PARSING
# ============================================
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠️ PyPDF2 not installed. PDF support disabled.")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("⚠️ python-docx not installed. Word document support disabled.")

try:
    import openpyxl
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False
    print("⚠️ openpyxl not installed. Excel support disabled.")


# ============================================
# INVOICE EXTRACTION FUNCTIONS
# ============================================
def extract_text_from_pdf(file_content: bytes) -> str:
    if not PDF_SUPPORT:
        return ""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_text_from_docx(file_content: bytes) -> str:
    if not DOCX_SUPPORT:
        return ""
    try:
        doc = Document(io.BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text for cell in row.cells if cell.text])
                if row_text:
                    text += "\n" + row_text
        return text
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def extract_text_from_xlsx(file_content: bytes) -> str:
    if not XLSX_SUPPORT:
        return ""
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
        text = ""
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"\n--- Sheet: {sheet_name} ---\n"
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join([str(cell) for cell in row if cell])
                if row_text:
                    text += row_text + "\n"
        return text
    except Exception as e:
        print(f"XLSX extraction error: {e}")
        return ""


def extract_invoice_data(text: str) -> dict:
    data = {}
    
    if not text or len(text.strip()) < 10:
        return {
            'vendor': 'Unknown',
            'total': 0.0,
            'tax': 0.0,
            'date': datetime.now().strftime('%Y-%m-%d'),
            'invoiceNumber': f'INV-{datetime.now().strftime("%Y%m%d%H%M%S")}'
        }
    
    vendor_patterns = [
        r'(?:Vendor|From|Company|Store|Merchant|Seller|Supplier)[:\s]+([^\n]+)',
        r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+Invoice',
        r'Bill To:?\s*([^\n]+)',
    ]
    for pattern in vendor_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            data['vendor'] = match.group(1).strip()[:100]
            break
    if 'vendor' not in data:
        data['vendor'] = 'Unknown'

    total_patterns = [
        r'(?:Total|Amount Due|Invoice Total|Grand Total|Balance Due)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)',
        r'(?:Total|Amount)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)',
        r'[\$£€]\s*([\d,]+\.?\d*)\s*(?:Total|Amount)',
    ]
    for pattern in total_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                data['total'] = float(match.group(1).replace(',', ''))
                break
            except:
                continue
    if 'total' not in data:
        data['total'] = 0.0

    tax_patterns = [
        r'(?:Tax|GST|VAT|HST)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)',
        r'(?:Sales Tax|Tax Amount)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)'
    ]
    for pattern in tax_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                data['tax'] = float(match.group(1).replace(',', ''))
                break
            except:
                continue
    if 'tax' not in data:
        data['tax'] = 0.0

    date_patterns = [
        r'(?:Date|Invoice Date|Issue Date|Created)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        r'(?:Date)[:\s]+(\d{4}-\d{2}-\d{2})',
        r'(\d{1,2}/\d{1,2}/\d{4})',
    ]
    for pattern in date_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            data['date'] = match.group(1)
            break
    if 'date' not in data:
        data['date'] = datetime.now().strftime('%Y-%m-%d')

    inv_patterns = [
        r'(?:Invoice|Invoice Number|INV|Bill|Receipt Number)[:\s#]+([A-Z0-9-]+)',
        r'Invoice\s*#?\s*([A-Z0-9-]+)',
        r'INV-\d+',
    ]
    for pattern in inv_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            data['invoiceNumber'] = match.group(1) if match.groups() else match.group(0)
            break
    if 'invoiceNumber' not in data:
        data['invoiceNumber'] = f'INV-{datetime.now().strftime("%Y%m%d%H%M%S")}'

    return data


# ============================================
# HEALTH & MODEL STATUS
# ============================================
@app.get("/")
def root():
    return {"message": "ProphetLedger API is running!", "status": "healthy", "version": "1.0.0"}

@app.get("/api/health")
def health():
    return {
        "status": "healthy",
        "services": {"supabase": supabase is not None, "groq": groq_client is not None},
        "environment": "production"
    }

@app.get("/api/models/status")
async def get_model_status():
    return hf_loader.get_model_status()


# ============================================
# INVOICE ENDPOINTS
# ============================================
@app.post("/api/invoices/extract-text")
async def extract_text_only(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user)
):
    try:
        contents = await file.read()
        file_type = file.content_type
        text = ""
        
        if file_type == 'application/pdf':
            text = extract_text_from_pdf(contents)
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            text = extract_text_from_docx(contents)
        elif file_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
            text = extract_text_from_xlsx(contents)
        else:
            text = contents.decode('utf-8', errors='ignore')
        
        if not text or len(text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Could not extract sufficient text from file")
        
        return {"text": text, "filename": file.filename, "file_type": file_type}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text: {str(e)}")


@app.post("/api/invoices/process")
async def process_invoice(
    file: UploadFile = File(...),
    extracted_text: str = Form(...),
    current_user = Depends(get_current_user)
):
    try:
        extracted_data = extract_invoice_data(extracted_text)
        extracted_data['file_name'] = file.filename
        extracted_data['file_type'] = file.content_type
        extracted_data['extraction_method'] = 'regex'
        return extracted_data
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process invoice: {str(e)}")


# ============================================
# AUTH ENDPOINTS
# ============================================
class LoginRequest(BaseModel):
    email: str
    password: str

class RegisterRequest(BaseModel):
    email: str
    full_name: str
    password: str

@app.post("/api/auth/register")
async def register(request: RegisterRequest):
    if not supabase:
        raise HTTPException(status_code=500, detail="Supabase not configured")
    try:
        response = supabase.auth.sign_up({
            "email": request.email,
            "password": request.password,
            "options": {"data": {"full_name": request.full_name}}
        })
        return {"message": "User created successfully", "user_id": response.user.id}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/auth/login")
async def login(request: LoginRequest):
    if not supabase:
        raise HTTPException(status_code=500, detail="Supabase not configured")
    try:
        response = supabase.auth.sign_in_with_password({
            "email": request.email,
            "password": request.password
        })
        return {
            "access_token": response.session.access_token,
            "token_type": "bearer",
            "user": {
                "id": response.user.id,
                "email": response.user.email,
                "full_name": response.user.user_metadata.get("full_name", ""),
                "role": "user",
                "mode_preference": "personal"
            }
        }
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid credentials")

@app.get("/api/auth/me")
async def get_current_user_route(current_user=Depends(get_current_user)):
    return {
        "id": current_user.id,
        "email": current_user.email,
        "full_name": current_user.user_metadata.get("full_name", ""),
        "is_active": True
    }


# ============================================
# CHATBOT ENDPOINT
# ============================================
class ChatRequest(BaseModel):
    query: str

@app.post("/api/chatbot/query")
async def chat(request: ChatRequest):
    if not groq_client:
        return {"query": request.query, "response": "Chatbot unavailable", "intent": "error", "confidence": 0}
    try:
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a helpful financial assistant."},
                {"role": "user", "content": request.query}
            ],
            temperature=0.7,
            max_tokens=500
        )
        return {"query": request.query, "response": completion.choices[0].message.content, "intent": "llm", "confidence": 0.95}
    except Exception as e:
        return {"query": request.query, "response": str(e), "intent": "error", "confidence": 0}


# ============================================
# RISK SCORE ENDPOINT (DYNAMIC)
# ============================================
@app.get("/api/dss/risk/score")
async def get_risk_score(request: Request):
    if not supabase:
        return {"risk_score": 50, "risk_level": "medium", "active_anomalies": 0, "recommendation": "Risk scoring unavailable"}
    
    auth_header = request.headers.get("Authorization")
    if not auth_header:
        return {"risk_score": 50, "risk_level": "medium", "active_anomalies": 0, "recommendation": "Please login"}
    
    token = auth_header.replace("Bearer ", "")
    try:
        user_data = supabase.auth.get_user(token)
        user_id = user_data.user.id
    except:
        return {"risk_score": 50, "risk_level": "medium", "active_anomalies": 0, "recommendation": "Invalid token"}
    
    try:
        anomalies_response = supabase.table("anomalies").select("id, anomaly_score, status").eq("user_id", user_id).execute()
        anomalies = anomalies_response.data or []
        pending_count = len([a for a in anomalies if a.get('status') == 'pending'])
        
        if pending_count == 0:
            risk_score = 25
            risk_level = "low"
            recommendation = "No anomalies detected. Good job!"
        elif pending_count <= 2:
            risk_score = 50
            risk_level = "medium"
            recommendation = f"Review {pending_count} pending anomaly(s)."
        else:
            risk_score = 75
            risk_level = "high"
            recommendation = f"Urgent: Review {pending_count} pending anomalies."
        
        return {"risk_score": risk_score, "risk_level": risk_level, "active_anomalies": pending_count, "recommendation": recommendation}
    except Exception as e:
        print(f"Error: {e}")
        return {"risk_score": 50, "risk_level": "medium", "active_anomalies": 0, "recommendation": "Unable to calculate"}


# ============================================
# KPI ENDPOINTS
# ============================================
@app.get("/api/dss/kpis")
async def get_kpis(mode: str = "personal"):
    return [
        {"id": 1, "title": "Financial Health", "value": 78, "change": 5.2, "trend": "up", "benchmark": 75, "status": "good", "recommendation": "Keep saving!"},
        {"id": 2, "title": "Cash Runway", "value": 12, "change": -2, "trend": "down", "benchmark": 12, "status": "warning", "recommendation": "Watch spending"},
        {"id": 3, "title": "Burn Rate", "value": 15000, "change": 8, "trend": "up", "benchmark": 10000, "status": "critical", "recommendation": "Cut expenses"},
        {"id": 4, "title": "Savings Rate", "value": 18, "change": 3, "trend": "up", "benchmark": 20, "status": "warning", "recommendation": "Save more"}
    ]


# ============================================
# FORECASTS
# ============================================
@app.get("/api/forecasts/trend/{metric}")
async def get_trend(metric: str, days: int = 90):
    dates = [(datetime.now() - timedelta(days=i)).strftime('%Y-%m-%d') for i in range(days, 0, -1)]
    base = 50000 if metric == "cashflow" else 32000
    values = np.cumsum(np.random.normal(100, 500, days)) + base
    history = [{"date": dates[i], "actual": round(float(values[i]), 2)} for i in range(len(dates))]
    last = values[-1]
    forecast_dates = [(datetime.now() + timedelta(days=i)).strftime('%Y-%m-%d') for i in range(1, 31)]
    forecast_values = np.linspace(last, last * 1.05, 30)
    for i in range(30):
        history.append({"date": forecast_dates[i], "actual": None, "forecast": round(float(forecast_values[i]), 2)})
    return {"metric": metric, "history": history, "anomalies": []}


# ============================================
# ANOMALIES
# ============================================
@app.get("/api/anomalies")
async def get_anomalies(limit: int = 10):
    return [{"id": 1, "date": "2024-05-15", "description": "Amazon Purchase", "amount": 1249.99, "category": "Shopping", "anomaly_score": 92, "status": "pending"}][:limit]


# ============================================
# TRANSACTIONS
# ============================================
@app.get("/api/transactions")
async def get_transactions(limit: int = 50):
    return [
        {"id": 1, "date": "2024-05-15", "description": "Starbucks", "amount": 5.75, "category": "Dining", "type": "expense"},
        {"id": 2, "date": "2024-05-14", "description": "Salary", "amount": 5000, "category": "Income", "type": "income"},
    ][:limit]


# ============================================
# CLASSIFICATION FALLBACK
# ============================================
class TransactionToClassify(BaseModel):
    description: str
    amount: float

@app.post("/api/transactions/classify")
async def classify_transaction(transaction: TransactionToClassify):
    keywords = {
        'Groceries': ['walmart', 'target', 'kroger', 'costco', 'aldi', 'whole foods'],
        'Dining': ['starbucks', 'mcdonalds', 'chipotle', 'restaurant', 'cafe', 'burger', 'pizza'],
        'Transport': ['uber', 'lyft', 'taxi', 'gas', 'shell', 'exxon', 'parking'],
        'Utilities': ['electric', 'water', 'internet', 'phone', 'comcast', 'att', 'verizon'],
        'Entertainment': ['netflix', 'spotify', 'disney', 'hulu', 'cinema', 'movie'],
        'Shopping': ['amazon', 'ebay', 'nike', 'adidas', 'clothing', 'shoes', 'best buy'],
        'Health': ['doctor', 'dental', 'hospital', 'pharmacy', 'gym'],
        'Rent': ['rent', 'apartment', 'lease', 'property'],
        'Income': ['salary', 'payroll', 'deposit', 'freelance', 'payment']
    }
    description_lower = transaction.description.lower()
    for category, words in keywords.items():
        if any(word in description_lower for word in words):
            return {"category": category, "confidence": 0.7, "method": "keyword"}
    if transaction.amount > 1000:
        return {"category": "Income", "confidence": 0.6, "method": "keyword"}
    return {"category": "Other", "confidence": 0.4, "method": "keyword"}