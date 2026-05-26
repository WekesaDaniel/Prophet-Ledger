# backend/app/api/invoices.py
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import Optional, List
from pydantic import BaseModel
from datetime import datetime
import re
import io
from app.database import get_db
from app.middleware.auth import get_current_user
from app.models.user import User
from app.models.invoice import Invoice

router = APIRouter(prefix="/api/invoices", tags=["Invoices"])

# Try to import optional dependencies
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("Warning: PyPDF2 not installed. PDF support disabled.")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not installed. Word document support disabled.")

try:
    import openpyxl
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False
    print("Warning: openpyxl not installed. Excel support disabled.")


class InvoiceResponse(BaseModel):
    id: int
    user_id: int
    vendor: Optional[str]
    total_amount: Optional[float]
    tax: Optional[float]
    date: Optional[datetime]
    pdf_url: Optional[str]
    invoice_number: Optional[str]
    extracted_data: Optional[dict]
    created_at: datetime
    
    class Config:
        from_attributes = True


class InvoiceCreate(BaseModel):
    vendor: str
    total_amount: float
    tax: float = 0
    date: datetime
    pdf_url: str
    invoice_number: str
    extracted_data: dict


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF using PyPDF2"""
    if not PDF_SUPPORT:
        return ""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from Word document"""
    if not DOCX_SUPPORT:
        return ""
    try:
        doc = Document(io.BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text for cell in row.cells if cell.text])
                if row_text:
                    text += "\n" + row_text
        return text
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def extract_text_from_xlsx(file_content: bytes) -> str:
    """Extract text from Excel file"""
    if not XLSX_SUPPORT:
        return ""
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
        text = ""
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"\n--- Sheet: {sheet_name} ---\n"
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join([str(cell) for cell in row if cell])
                if row_text:
                    text += row_text + "\n"
        return text
    except Exception as e:
        print(f"XLSX extraction error: {e}")
        return ""


def extract_invoice_data(text: str) -> dict:
    """Extract invoice data using regex patterns"""
    data = {}
    
    if not text:
        return {
            'vendor': 'Unknown',
            'total': 0.0,
            'tax': 0.0,
            'date': datetime.now().strftime('%Y-%m-%d'),
            'invoiceNumber': f'INV-{datetime.now().strftime("%Y%m%d%H%M%S")}'
        }
    
    # Extract vendor name
    vendor_patterns = [
        r'(?:Vendor|From|Company|Store|Merchant|Seller|Supplier)[:\s]+([^\n]+)',
        r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+Invoice',
        r'Bill To:?\s*([^\n]+)',
        r'(?:Sold by|Provided by)[:\s]+([^\n]+)'
    ]
    for pattern in vendor_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            data['vendor'] = match.group(1).strip()[:100]
            break
    if 'vendor' not in data:
        data['vendor'] = 'Unknown'

    # Extract total amount
    total_patterns = [
        r'(?:Total|Amount Due|Invoice Total|Grand Total|Balance Due)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)',
        r'(?:Total|Amount)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)',
        r'[\$£€]\s*([\d,]+\.?\d*)\s*(?:Total|Amount)',
        r'(?:Subtotal|Sub total)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)'
    ]
    for pattern in total_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                data['total'] = float(match.group(1).replace(',', ''))
                break
            except:
                continue
    if 'total' not in data:
        data['total'] = 0.0

    # Extract tax
    tax_patterns = [
        r'(?:Tax|GST|VAT|HST)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)',
        r'(?:Tax|GST|VAT).*?[\$£€]\s*([\d,]+\.?\d*)',
        r'(?:Sales Tax|Tax Amount)[:\s]*[\$£€]?\s*([\d,]+\.?\d*)'
    ]
    for pattern in tax_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                data['tax'] = float(match.group(1).replace(',', ''))
                break
            except:
                continue
    if 'tax' not in data:
        data['tax'] = 0.0

    # Extract date
    date_patterns = [
        r'(?:Date|Invoice Date|Issue Date|Created)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        r'(?:Date)[:\s]+(\d{4}-\d{2}-\d{2})',
        r'(\d{1,2}/\d{1,2}/\d{4})',
        r'(\d{4}-\d{2}-\d{2})'
    ]
    for pattern in date_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            data['date'] = match.group(1)
            break
    if 'date' not in data:
        data['date'] = datetime.now().strftime('%Y-%m-%d')

    # Extract invoice number
    inv_patterns = [
        r'(?:Invoice|Invoice Number|INV|Bill|Receipt Number)[:\s#]+([A-Z0-9-]+)',
        r'Invoice\s*#?\s*([A-Z0-9-]+)',
        r'INV-\d+',
        r'#\s*([A-Z0-9-]+)'
    ]
    for pattern in inv_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            data['invoiceNumber'] = match.group(1) if match.groups() else match.group(0)
            break
    if 'invoiceNumber' not in data:
        data['invoiceNumber'] = f'INV-{datetime.now().strftime("%Y%m%d%H%M%S")}'

    return data


@router.post("/extract-text")
async def extract_text_only(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Extract raw text from uploaded file (no OCR processing)"""
    try:
        contents = await file.read()
        file_type = file.content_type
        text = ""
        
        if file_type == 'application/pdf':
            text = extract_text_from_pdf(contents)
        elif file_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            text = extract_text_from_docx(contents)
        elif file_type in [
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel'
        ]:
            text = extract_text_from_xlsx(contents)
        else:
            raise HTTPException(status_code=400, detail=f"File type not supported for text extraction: {file_type}")
        
        return {"text": text, "filename": file.filename, "file_type": file_type}
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text: {str(e)}")


@router.post("/process")
async def process_invoice(
    file: UploadFile = File(...),
    extracted_text: str = Form(...),
    current_user: User = Depends(get_current_user)
):
    """Process extracted text and return structured invoice data"""
    try:
        # Extract invoice data using regex
        extracted_data = extract_invoice_data(extracted_text)
        
        # Add metadata
        extracted_data['file_name'] = file.filename
        extracted_data['file_type'] = file.content_type
        extracted_data['extraction_method'] = 'regex'
        extracted_data['text_length'] = len(extracted_text)
        
        print(f"Successfully extracted: Vendor={extracted_data['vendor']}, Total={extracted_data['total']}")
        
        return extracted_data
        
    except Exception as e:
        print(f"Processing error: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Failed to process invoice data: {str(e)}")


@router.get("/", response_model=List[InvoiceResponse])
def get_invoices(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all invoices for current user"""
    invoices = db.query(Invoice).filter(Invoice.user_id == current_user.id).order_by(Invoice.created_at.desc()).all()
    return invoices


@router.get("/{invoice_id}", response_model=InvoiceResponse)
def get_invoice(
    invoice_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific invoice"""
    invoice = db.query(Invoice).filter(
        Invoice.id == invoice_id,
        Invoice.user_id == current_user.id
    ).first()
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    return invoice


@router.post("/", response_model=InvoiceResponse)
def create_invoice(
    invoice: InvoiceCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new invoice record"""
    db_invoice = Invoice(
        user_id=current_user.id,
        vendor=invoice.vendor,
        total_amount=invoice.total_amount,
        tax=invoice.tax,
        date=invoice.date,
        pdf_url=invoice.pdf_url,
        invoice_number=invoice.invoice_number,
        extracted_data=invoice.extracted_data
    )
    db.add(db_invoice)
    db.commit()
    db.refresh(db_invoice)
    return db_invoice


@router.delete("/{invoice_id}")
def delete_invoice(
    invoice_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete an invoice"""
    invoice = db.query(Invoice).filter(
        Invoice.id == invoice_id,
        Invoice.user_id == current_user.id
    ).first()
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    db.delete(invoice)
    db.commit()
    return {"message": "Invoice deleted successfully"}